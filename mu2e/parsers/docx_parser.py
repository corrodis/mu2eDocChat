"""
DOCX (Word document) parser
"""

import io
from PIL import Image
from .base_parser import BaseParser

class DOCXParser(BaseParser):
    """Parser for DOCX files"""
    
    def get_text(self, rescale_image_max_dim=500):
        """Extract text and images from DOCX document"""
        try:
            import docx
            
            doc = docx.Document(self.doc)
            text_parts = []
            images = []
            image_cnt = 0
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract images from document relationships
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        image_data = rel.target_part.blob
                        img = Image.open(io.BytesIO(image_data))
                        img = self._resize_image(img, rescale_image_max_dim)
                        img_base64 = self._image_to_base64(img, img.format)
                        
                        images.append(img_base64)
                        text_parts.append(f"[Image {image_cnt}]")
                        image_cnt += 1
                    except Exception as e:
                        print(f"Error processing image: {e}")
            
            # Extract text from tables
            for table in doc.tables:
                table_text = "\n**Table:**\n"
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    table_text += f"| {row_text} |\n"
                text_parts.append(table_text)
            
            full_text = '\n'.join(text_parts)
            cleaned_text = self._clean_text(full_text)
            
            return cleaned_text, images
            
        except ImportError:
            print("python-docx not installed. Install with: pip install python-docx")
            return "", []
        except Exception as e:
            print(f"Error parsing DOCX document: {e}")
            return "", []